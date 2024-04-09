from fpdf import FPDF
import textwrap
import os
import aspose.words as aw
from docx import Document
import re

#Tieng anh
def TXTtoPDF(inputtxt,fname):
    file = open(inputtxt)
    text = file.read()
    a4_width_mm = 210
    pt_to_mm = 0.35
    fontsize_pt = 10
    fontsize_mm = fontsize_pt * pt_to_mm
    margin_bottom_mm = 10
    character_width_mm = 7 * pt_to_mm
    width_text = a4_width_mm/character_width_mm

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.set_auto_page_break(True, margin=margin_bottom_mm)
    pdf.add_page()
    pdf.set_font("Courier", size=fontsize_pt)
    splitted = text.split('\n')

    for line in splitted:
        lines = textwrap.wrap(line, width_text)

        if len(lines) == 0:
            pdf.ln()

        for wrap in lines:
            pdf.cell(0, fontsize_mm, wrap, ln=1)

    completeName=fname
    pdf.output(completeName, 'F')

def TXTtoWORD(inputtxt,dest,fname):
    file = open(inputtxt)
    text = file.read()
    text = re.sub(r'[^\x00-\x7F]+|\x0c',' ', text)
    document = Document()
    document.add_paragraph(text)
    completeName=os.path.join(dest, fname+".docx")  
    document.save(completeName)

#ca 2 deu dung duoc cho tieng anh va tieng viet khong loi
def TXTtoPDFwithWATERMARK(inputtxt,dest,fname):
    completeName=os.path.join(dest, fname+".pdf") 
    pdf=aw.Document(inputtxt)
    pdf.save(completeName) 

def TXTtoDOCwithWATERMARK(inputtxt,dest,fname):
    completeName=os.path.join(dest, fname+".docx") 
    doc=aw.Document(inputtxt)
    doc.save(completeName) 

# TXTtoPDFwithWATERMARK('./data/change.txt','./data/','OUTPUT')

# TXTtoDOCwithWATERMARK('./data/change.txt','./data/','OUTPUT')
# ver2('./data/change.txt','./data/','OUTPUT')

# TXTtoPDFwithWATERMARK('./data/vn.txt','./data/','VN')

# TXTtoPDF('./data/change.txt','./data/','OUTPUT')
# TXTtoWORD('./data/change.txt','./data/','OUTPUT')